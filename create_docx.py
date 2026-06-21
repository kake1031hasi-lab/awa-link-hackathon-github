import os
import glob
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    doc = Document()

    # タイトルスタイル設定
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("AWA-LINK ハッカソン向け「Google Cloud ドレスコード」提案書")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(22)
    title_run.bold = True

    # サブタイトル
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run("次世代エージェントとGoogle Cloudインフラによる臨床支援AI")
    sub_run.font.name = 'Arial'
    sub_run.font.size = Pt(14)
    sub_run.italic = True

    doc.add_paragraph("\n")

    # セクション1
    doc.add_heading("1. 提案する新・システム構成図（Google Cloud構成）", level=1)
    doc.add_paragraph(
        "現在の Cloudflare Workers 構成から、Google Cloud のフルマネージドサービスと新しいAI標準プロトコル（MCP）を組み込んだ次世代エージェント構成に「お色直し」した設計図です。"
    )
    
    # 構成図の説明
    doc.add_paragraph(
        "■ アーキテクチャ概要:\n"
        "1. LINEユーザーからの Webhook (POST) を Google Cloud Run（サーバーレスコンテナ実行環境）で受取・処理します。\n"
        "2. Gemini APIとの接続には Interactions API を利用し、セッション間の文脈を保持（ステートフル対話）します。\n"
        "3. スプレッドシートや BigQuery などの Google サービスとの連携には、標準プロトコルである Model Context Protocol (MCP) を経由して接続します。"
    )

    doc.add_paragraph("\n")

    # セクション2
    doc.add_heading("2. ハッカソンで審査員を「刺す」4つのアピールポイント", level=1)
    
    # アピール1
    p1 = doc.add_paragraph()
    r1 = p1.add_run("🌟 アピール①：Agents CLI & Cloud Run による DevOps パイプライン\n")
    r1.bold = True
    p1.add_run(
        "「今回のデプロイには、Google公式の Agents CLI (google-agents-cli) と Cloud Run を採用しました。"
        "infra コマンドで自動的に Google Cloud 上にセキュアなコンテナ環境を構築し、deploy コマンドで一発出荷する"
        "先進的な DevOps パイプラインを構築しています。これにより、インフラの運用コストを最小化しつつ、エンタープライズ品質のデリバリー速度を実現しました。」"
    )

    # アピール2
    p2 = doc.add_paragraph()
    r2 = p2.add_run("🌟 アピール②：Interactions API（次世代の標準）による文脈保持\n")
    r2.bold = True
    p2.add_run(
        "「Geminiの呼び出しには、従来のステートレスな generateContent ではなく、最新の Interactions API を活用しています。"
        "会話セッション（文脈）をサーバー側でステートフルに維持・管理するアーキテクチャにより、臨床での『長時間のブレインストーミングや壁打ち』でも、"
        "途中で文脈や前提条件を見失わない『対話の命綱』を実装しました。」"
    )

    # アピール3
    p3 = doc.add_paragraph()
    r3 = p3.add_run("🌟 アピール③：Google公式「MCP（Model Context Protocol）」による外部連携\n")
    r3.bold = True
    p3.add_run(
        "「API連携をコードにベタ書きする時代は終わりました。本システムでは、Google公式の MCP Hub を採用し、"
        "AIが安全かつ動的に外部リソースにアクセスできる『手足』を構築しています。スプレッドシートへのログ記録、"
        "BigQueryへのデータ転送、Googleカレンダーへの連携はすべて標準化された MCP 経由で動作（または構想）しています。」"
    )

    doc.add_paragraph("\n")

    # セクション3
    doc.add_heading("3. プレゼン・デモを最高に盛り上げる「ハリボテ（完成イメージ）」", level=1)
    
    # 画像の取得
    brain_dir = "C:/Users/kake1/.gemini/antigravity-ide/brain/fcc05613-6fb3-4199-8b54-38832d1f7799"
    dashboard_files = glob.glob(os.path.join(brain_dir, "dashboard_mockup_*.png"))
    lecturer_files = glob.glob(os.path.join(brain_dir, "lecturer_dispatch_mockup_*.png"))

    doc.add_heading("デモ①：BigQueryを用いた「地域スタッフの不満・不安」リアルタイム解析ダッシュボード", level=2)
    doc.add_paragraph(
        "LINEから流れてくる会話ログ（ElasticsearchやBigQueryに保存されたログ）を元に、"
        "どのリハビリ現場で・どのような論文情報が不足しているか、スタッフが何に不安を感じているかを時系列で自動分析するダッシュボードのモックアップです。"
    )
    if dashboard_files:
        doc.add_picture(dashboard_files[0], width=Inches(6.0))
        doc.add_paragraph("（図1. リアルタイム時系列解析ダッシュボード画面イメージ）")
    else:
        doc.add_paragraph("【画像が見つかりませんでした】")

    doc.add_paragraph("\n")

    doc.add_heading("デモ②：スプレッドシート連動「悩み合わせた最適な講師の自動派遣提案」", level=2)
    doc.add_paragraph(
        "スタッフの臨床の悩みをAIが分析し、スプレッドシート上で「どの分野の講師をどのステーションに派遣するのが最適か」を自動でマッチング・提案するデモ画面です。"
    )
    if lecturer_files:
        doc.add_picture(lecturer_files[0], width=Inches(6.0))
        doc.add_paragraph("（図2. 講師自動派遣マッチング画面イメージ）")
    else:
        doc.add_paragraph("【画像が見つかりませんでした】")

    doc.add_paragraph("\n")

    # セクション4
    doc.add_heading("4. Google Cloudへの移行・実現に必要なステップ", level=1)
    doc.add_paragraph(
        "もし「実際に Google Cloud 上で動作するもの」に切り替えたい場合、以下のステップでコードの移行を行います。\n\n"
        "1. Dockerfile の追加：Cloudflare Workers の Node.js コードを Cloud Run（コンテナ）で動かすため、軽量な Dockerfile を作成します。\n"
        "2. ルーティングとライブラリの調整：Cloudflare の fetch イベントリスナーから、通常の Express.js（Node.js）サーバーのコードに書き換えます。\n"
        "3. gcloud / Agents CLI によるデプロイ：gcloud run deploy コマンドで Google Cloud 上に数分でデプロイ可能です。"
    )

    # 保存
    output_path = "C:/Users/kake1/.gemini/antigravity-ide/scratch/awa-link/AWA-LINK_Proposal.docx"
    doc.save(output_path)
    print(f"Document saved to {output_path}")

if __name__ == "__main__":
    main()
